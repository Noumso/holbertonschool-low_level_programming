from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH doc
= Document()
# Titre stylisé
title = doc.add_heading(&#39;PLAN D’ACTIONS GPEC 2025-2027&#39;, level=0) title.alignment =
WD_ALIGN_PARAGRAPH.CENTER
run = title.runs[^0] run.font.size = Pt(28)
run.font.color.rgb = RGBColor(0, 32, 96) # Bleu institutionnel
# Sous-titre
subtitle = doc.add_paragraph(&#39;AISTHESIS FORMATION&#39;)
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle.runs[^0].italic = True
# Bloc métadonnées
meta = doc.add_paragraph()
meta.add_run(&#39;Date : Avril 2025\n&#39;).bold = True meta.add_run(&#39;Auteur :
Chargé de Développement RH\n&#39;)
meta.add_run(&#39;Référence : GPEC-MEDSOC-2025&#39;).font.color.rgb = RGBColor(128, 128, 128)